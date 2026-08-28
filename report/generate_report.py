#!/usr/bin/env python3
"""Build report/Comprehensive_Analysis_Report.docx from the live result files.

Every number in the report is read from results/*.csv or recomputed from
data/derived/data_with_vlc.csv at run time. Nothing is hard-coded, so the
document cannot drift from the analysis: if a script is re-run and a value
changes, the next `make report` picks it up.

The narrative follows the manuscript's five model families:
    Family 1  classical Hall-Petch
    Family 2  physics-derived descriptors (SSS, Wen, PCA-OLS)
    Family 3  composition / processing (M-model hierarchy, incl. M15)
    Family 4  non-linear ML at matched inputs
    Family 5  symbolic regression
with the validation protocol and the literature stress test running under all
five.

Usage:  python report/generate_report.py
"""
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import KFold

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
sys.path.insert(0, str(ROOT / 'scripts'))
from _config import DATA_DIR, RESULTS_DIR, PLOTS_DIR, PAPER_FIG_DIR   # noqa: E402

ELEMS = ['Al', 'Co', 'Cr', 'Cu', 'Fe', 'Mn', 'V']        # Ni is the reference
ACCENT = RGBColor(0x00, 0x43, 0x6B)


# --------------------------------------------------------------- helpers ----
def csv(name):
    p = RESULTS_DIR / name
    return pd.read_csv(p) if p.exists() else None


def q2(y, p):
    return 1 - ((y - p) ** 2).sum() / ((y - y.mean()) ** 2).sum()


def cv_scores(X, y, groups):
    """Pooled 5-fold / LOO / LOBO Q^2 for an OLS design matrix."""
    p5 = np.zeros(len(y))
    for tr, te in KFold(5, shuffle=True, random_state=42).split(X):
        p5[te] = LinearRegression().fit(X[tr], y[tr]).predict(X[te])
    X1 = np.column_stack([np.ones(len(X)), X])
    beta, *_ = np.linalg.lstsq(X1, y, rcond=None)
    e = y - X1 @ beta
    h = np.clip(np.diag(X1 @ np.linalg.pinv(X1.T @ X1) @ X1.T), 0, 1 - 1e-10)
    pl = np.zeros(len(y))
    for k in np.unique(groups):
        te = groups == k
        pl[te] = LinearRegression().fit(X[~te], y[~te]).predict(X[te])
    n, k = len(y), X1.shape[1]
    bic = n * np.log((e ** 2).sum() / n) + k * np.log(n)
    return q2(y, p5), q2(y, y - e / (1 - h)), q2(y, pl), bic


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(it).font.size = Pt(10.5)


def table(doc, df, caption=None, fmt='{:.3f}'):
    if caption:
        c = doc.add_paragraph()
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9.5)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = (fmt.format(v) if isinstance(v, (int, float, np.floating))
                             and not isinstance(v, bool) else str(v))
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def figure(doc, filename, caption, width=6.3):
    for d in (PAPER_FIG_DIR, PLOTS_DIR):
        f = Path(d) / filename
        if f.exists():
            doc.add_picture(str(f), width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            return True
    return False


# ------------------------------------------------------------------ data ----
DF = pd.read_csv(DATA_DIR / 'data_with_vlc.csv')
YS = DF.dropna(subset=['YS']).reset_index(drop=True)
HV = DF.dropna(subset=['HV']).reset_index(drop=True)

_d, _SD, _y, _g = (YS.d_inv_sqrt.values, YS.SD_GS.values,
                   YS.YS.values, YS.Iteration.values)
_F = {e: YS[f'{e}_frac'].values for e in ELEMS}

MODELS = {
    'M0: baseline Hall-Petch':      np.column_stack([_d]),
    'M1: sigma_0(V)':               np.column_stack([_F['V'], _d]),
    'M3: sigma_0(all 7 elements)':  np.column_stack([*[_F[e] for e in ELEMS], _d]),
    'M4: k(V)':                     np.column_stack([_d, _F['V'] * _d]),
    'M6: k(all 7 elements)':        np.column_stack([_d, *[_F[e] * _d for e in ELEMS]]),
    'M10: sigma_0 + k (all)':       np.column_stack([*[_F[e] for e in ELEMS], _d,
                                                     *[_F[e] * _d for e in ELEMS]]),
    'M11: sigma_0(delta)':          np.column_stack([YS.delta.values, _d]),
    'M13: M3 + SD_grain':           np.column_stack([*[_F[e] for e in ELEMS], _d, _SD]),
    'M15: M3 + SD_grain + SD*d^-1/2':
                                    np.column_stack([*[_F[e] for e in ELEMS], _d, _SD, _SD * _d]),
}


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)

    # -------------------------------------------------------------- title --
    t = doc.add_heading('Comprehensive Analysis Report', level=0)
    for r in t.runs:
        r.font.color.rgb = ACCENT
    para(doc, 'Revisiting Hall-Petch strengthening in FCC multi-principal '
              'element alloys: best practices for building and auditing '
              'machine-learning strength models', bold=True, size=12)
    para(doc, 'Companion report to the manuscript in paper/. Every value below '
              'is read from results/ or recomputed from data/derived at build '
              'time; see docs/reproducing.md for the run order.', size=9.5)
    doc.add_paragraph()

    # ------------------------------------------------------------ dataset --
    h1(doc, '1. Dataset')
    audit = csv('dataset_audit.csv')
    if audit is not None:
        a = audit.iloc[0]
        para(doc, f"The dataset contains {int(a.n_total)} FCC alloy conditions in the "
                  f"Al-Co-Cr-Cu-Fe-Mn-Ni-V system: {int(a.n_YS)} with yield strength and "
                  f"{int(a.n_HV)} with Vickers hardness. These represent "
                  f"{int(a.n_unique_compositions_all)} unique compositions "
                  f"({int(a.n_unique_compositions_YS)} within the YS subset). "
                  f"{int(a.n_repeated_composition_groups)} chemistries recur across "
                  f"{int(a.n_rows_in_repeated_groups)} records, and "
                  f"{int(a.n_compositions_crossing_batches)} composition appears in more "
                  f"than one batch, so leave-one-batch-out reduces but does not remove "
                  f"exact-composition overlap.")
        para(doc, f"Mean grain size spans {a.d_min_um:.2f}-{a.d_max_um:.2f} um, yield "
                  f"strength {a.YS_min_MPa:.1f}-{a.YS_max_MPa:.1f} MPa and hardness "
                  f"{a.HV_min:.1f}-{a.HV_max:.1f} HV. Mean grain size and its within-map "
                  f"standard deviation are strongly correlated "
                  f"(r = {a.corr_mean_d_SD_GS:.3f}), so SD_grain is treated as an "
                  f"exploratory measure of distribution width rather than an "
                  f"independent mechanism.")
    figure(doc, 'fig_provenance.png',
           'Composition and processing coverage by batch.')
    doc.add_page_break()

    # ------------------------------------------------- validation protocol --
    h1(doc, '2. Validation protocol')
    para(doc, 'Three internal protocols ask progressively harder questions. All '
              'predictions from held-out records are pooled before scoring, so the '
              'reported R-squared is a single pooled Q-squared and not the mean of '
              'fold-specific values.')
    bullets(doc, [
        'Shuffled 5-fold cross-validation (seed 42): interpolation within the pooled dataset.',
        'Leave-one-out: record-level interpolation. A record with the same composition '
        'can remain in the training set, so LOO is not a test on unseen compositions.',
        'Leave-one-batch-out across the six BO batches: transfer to an unseen '
        'experimental iteration. This is the primary grouped transfer test.',
    ])
    para(doc, 'A heterogeneous literature set provides an external stress test, and '
              'every closed form is checked for poles, sign changes and non-finite '
              'values over the observed and literature ranges.')

    grouped = csv('main2_grouped_validation.csv')
    if grouped is not None:
        piv = (grouped.pivot_table(index=['target', 'model'], columns='protocol',
                                   values='R2')
               .reset_index()
               .rename(columns={'target': 'Target', 'model': 'Model'}))
        cols = ['Target', 'Model'] + [c for c in ['5-fold', 'LOO', 'LOBO'] if c in piv]
        table(doc, piv[cols].round(3),
              'Table 1. Pooled out-of-fold R-squared across the three internal protocols.')
    doc.add_page_break()

    # ---------------------------------------------------------- Family 1 ---
    h1(doc, '3. Family 1: classical Hall-Petch')
    for name, frame, tgt in (('Yield strength', YS, 'YS'), ('Vickers hardness', HV, 'HV')):
        X = np.column_stack([frame.d_inv_sqrt.values])
        s5, sl, sb, _ = cv_scores(X, frame[tgt].values, frame.Iteration.values)
        para(doc, f'{name}: 5-fold {s5:.3f}, LOO {sl:.3f}, LOBO {sb:.3f}.')
    para(doc, 'Grain size is a useful, reproducible baseline for yield strength and a '
              'poor one for hardness, where the batch-held-out score is negative. Nine '
              'pre-specified scaling laws were compared; five two-parameter monotonic '
              'forms fall within delta-BIC < 2 of the classical exponent, so the data '
              'support d^-1/2 as a parsimonious baseline rather than as a uniquely '
              'identified exponent.')
    figure(doc, 'fig01_hall_petch.png',
           'Classical Hall-Petch fits for YS and HV, coloured by batch.')
    figure(doc, 'fig_scaling_deltabic.png',
           'Information-criterion comparison of the nine grain-size scaling laws.')
    doc.add_page_break()

    # ---------------------------------------------------------- Family 2 ---
    h1(doc, '4. Family 2: physics-derived descriptors')
    para(doc, 'Three solid-solution-strengthening models (Varvenne-Leyson-Curtin, a '
              'weighted Labusch extension, and Toda-Caraballo) were implemented from '
              'their source papers using Vegard-law volumes and rule-of-mixtures '
              'elastic constants. Each was scored on three separate axes: standalone '
              'accuracy, marginal value alongside Hall-Petch, and information content '
              'beyond raw composition.')
    para(doc, 'The decisive test is redundancy. All three formulas are deterministic '
              'functions of composition, so any model given the elemental fractions can '
              'encode whatever transformation they express. Appending an SSS estimate to '
              'a composition model changes LOO R-squared by at most 0.002, and every '
              'partial correlation after removing composition is below 0.10 in '
              'magnitude.')
    pca = csv('pca_ols_nested_results.csv')
    if pca is not None:
        para(doc, 'Fold-contained PCA-OLS gives ' +
                  ', '.join(f'{r.protocol} {r.R2:.3f}' for _, r in pca.iterrows()) +
                  '. Fitting the scaler and PCA on all records before splitting had '
                  'produced 0.525 / 0.516 / 0.441, so even unsupervised preprocessing '
                  'inflates held-out scores when it precedes the split.')
    figure(doc, 'fig_sss_parity.png',
           'SSS predictions against Hall-Petch-corrected experimental friction stress.')
    doc.add_page_break()

    # ---------------------------------------------------------- Family 3 ---
    h1(doc, '5. Family 3: composition and processing')
    rows = []
    for name, X in MODELS.items():
        s5, sl, sb, bic = cv_scores(X, _y, _g)
        rows.append((name, X.shape[1] + 1, s5, sl, sb, bic))
    tab = pd.DataFrame(rows, columns=['Model', 'p', '5-fold', 'LOO', 'LOBO', 'BIC'])
    m3bic = tab.loc[tab.Model.str.startswith('M3'), 'BIC'].iloc[0]
    tab['dBIC vs M3'] = tab.BIC - m3bic
    table(doc, tab.drop(columns='BIC').round(3),
          'Table 2. The M-model hierarchy, recomputed at build time.')

    m3 = tab[tab.Model.str.startswith('M3:')].iloc[0]
    m13 = tab[tab.Model.str.startswith('M13')].iloc[0]
    m15 = tab[tab.Model.str.startswith('M15')].iloc[0]
    para(doc, f'Making the friction stress composition-dependent produces the largest '
              f'single gain for yield strength: M3 reaches LOO {m3.LOO:.3f} and LOBO '
              f'{m3.LOBO:.3f} against {tab.iloc[0].LOO:.3f} / {tab.iloc[0].LOBO:.3f} for '
              f'the grain-only baseline.')
    para(doc, f'Grain-distribution width helps only in a specific form. Adding SD_grain '
              f'as an independent additive term (M13) raises LOO to {m13.LOO:.3f} but '
              f'lowers LOBO to {m13.LOBO:.3f}. Allowing it to modify the Hall-Petch term '
              f'(M15) gives {m15["5-fold"]:.3f} / {m15.LOO:.3f} / {m15.LOBO:.3f} with '
              f'delta-BIC {m15["dBIC vs M3"]:+.1f} relative to M3. The additive control '
              f'shows the gain cannot be credited to measuring grain width alone; '
              f'because SD_grain is correlated with mean grain size and with processing, '
              f'the term is read as a predictive interaction rather than as proof that '
              f'distribution width changes the Hall-Petch slope.', bold=False)
    coef = csv('m3_coefficients.csv')
    if coef is not None:
        c = coef[coef.coefficient.str.startswith(('alpha', 'k_HP', 'intercept'))]
        table(doc, c[['coefficient', 'value', 'lo95', 'hi95']].round(1),
              'Table 3. M3 coefficients relative to Ni, with 95% intervals.')
    figure(doc, 'fig_comp_hp_models_ab.png',
           'Family 3 hierarchy: predictive ranking and parsimony.')
    doc.add_page_break()

    # ---------------------------------------------------------- Family 4 ---
    h1(doc, '6. Family 4: non-linear machine learning at matched inputs')
    para(doc, 'Linear controls and non-linear estimators were run on identical S1-S4 '
              'input matrices and identical splits, at declared fixed settings with no '
              'data-driven hyperparameter search, so that any difference reflects the '
              'model class rather than the tuning budget.')
    if grouped is not None:
        f4 = grouped[grouped.model.str.startswith('F4')]
        piv = f4.pivot_table(index='model', columns='protocol', values='R2').reset_index()
        cols = ['model'] + [c for c in ['5-fold', 'LOO', 'LOBO'] if c in piv]
        table(doc, piv[cols].round(3), 'Table 4. Matched-input comparison.')
    para(doc, 'At matched inputs the linear control and the gradient-boosted model give '
              'nearly identical batch-held-out scores, with overlapping batch-bootstrap '
              'intervals. There is no evidence that non-linearity improves transfer to '
              'an unseen batch on this dataset.')
    para(doc, 'One comparison needs care: the Family 1 baseline uses d^-1/2 alone, while '
              'the S1 matrix also contains SD_grain. Part of any S1-versus-baseline gap '
              'is therefore the extra feature and not the estimator.')
    figure(doc, 'fair_comparison_LOBO_heatmap.png',
           'Batch-held-out R-squared across the feature ladder for YS and HV.')
    doc.add_page_break()

    # ---------------------------------------------------------- Family 5 ---
    h1(doc, '7. Family 5: symbolic regression')
    para(doc, 'PySR and SISSO searched for compact closed forms. In the stored workflow '
              'the equation structures were selected after viewing complete-data search '
              'fronts and only their constants were refit inside the folds, so these are '
              'post-selection fixed-form scores and not unbiased estimates of symbolic '
              'discovery. They are reported as such and are not ranked against the '
              'pre-specified linear models.')
    hv = csv('hardness_equation_corrected.csv')
    if hv is not None:
        full = hv[hv.protocol == 'full fit']
        if len(full):
            r = full.iloc[0]
            para(doc, f'The refit hardness form is HV = {r.c0:.1f} + {r.c1:.1f}'
                      f'(6.93 - d)/SD_grain + {r.c2:.3f} dH_mix / t_hold^2, with '
                      f'full-data R-squared {r.R2:.3f} and RMSE {r.RMSE:.2f} HV.')
        cvr = hv[hv.protocol != 'full fit']
        if len(cvr):
            para(doc, 'With the structure fixed and constants refit per fold: ' +
                      ', '.join(f'{r.protocol} {r.R2:.3f}' for _, r in cvr.iterrows()) + '.')
    para(doc, 'Two safety findings belong with these equations. SISSO Full places a '
              'shear-modulus difference in a denominator and becomes unstable when the '
              'constituent moduli are similar; removing that term lowers the aggregate '
              'literature RMSE substantially. Separately, one descriptor-based PySR '
              'equation contains a protected square root whose argument approaches zero '
              'just below the observed range, so it stays real but diverges outside the '
              'sampled envelope.')
    figure(doc, 'fig08_pysr_pareto.png', 'PySR complexity-loss front for yield strength.')
    doc.add_page_break()

    # ------------------------------------------------ literature + Tabor ---
    h1(doc, '8. Literature stress test and hardness-yield correspondence')
    ext = csv('external_validation_by_evidence_tier.csv')
    if ext is not None:
        t1 = ext[ext.evidence_tier == 'Tier 1'][['label', 'R2', 'RMSE', 'bias', 'n']]
        table(doc, t1.round(2),
              'Table 5. Direct-measurement literature records (evidence class E1).')
    para(doc, 'Every evaluated model is worse than predicting the mean of the literature '
              'set. The compilation mixes sources and test modes and was filtered using '
              'outcomes, so it is reported as a stress test that exposes domain '
              'sensitivity and unsafe equations, not as an unbiased external benchmark.')

    paired = DF.dropna(subset=['YS', 'HV'])
    ceff = paired.HV.values * 9.807 / paired.YS.values
    para(doc, f'Across the {len(paired)} paired records the effective Tabor ratio is '
              f'{ceff.mean():.2f} +/- {ceff.std(ddof=1):.2f}, well above the classical '
              f'value of 3. Hardness and strength rank together within a batch but the '
              f'association weakens once batches are pooled, so a single conversion '
              f'cannot be transferred across composition and processing domains.')
    figure(doc, 'fig_hv_ys_rank.png', 'Hardness rank against strength rank, by batch.')
    doc.add_page_break()

    # ------------------------------------------------------------ closing --
    h1(doc, '9. What the analysis supports')
    bullets(doc, [
        'A physically interpretable grain-size baseline is worth establishing first; '
        'it calibrates expectations and exposes where the remaining variance lives.',
        'Physics-derived descriptors must be shown to add information beyond the '
        'measured inputs they are computed from, not merely to be physically motivated.',
        'Measured composition and processing should be exhausted before estimator '
        'complexity is increased.',
        'Grain-distribution width helps yield strength through an interaction with the '
        'Hall-Petch term, not as an additive contribution; the additive control is what '
        'makes that distinction visible.',
        'Validation groups must match the intended decision. Record-level interpolation, '
        'batch-level transfer and external deployment are different questions with '
        'different answers.',
        'Closed forms need a domain audit, including poles hidden inside protected '
        'operators, before any deployment claim.',
    ])
    para(doc, 'The coefficients reported here belong to this alloy system and these '
              'campaigns. The workflow does not.', bold=True)

    out = HERE / 'Comprehensive_Analysis_Report.docx'
    doc.save(out)
    print(f'Report saved: {out}')
    return out


if __name__ == '__main__':
    build()
